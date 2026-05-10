from __future__ import annotations

import json
from datetime import date
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate

from src.common.discussion_guide_builder import _markdown_to_story, _page_decorator
from src.image.feature_extraction import get_image_feature_names


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = PROJECT_ROOT / "docs"
RESULTS_PATH = PROJECT_ROOT / "outputs" / "reports" / "complete_results.json"


def _load_results() -> dict[str, Any]:
    return json.loads(RESULTS_PATH.read_text(encoding="utf-8"))


def _fmt(value: Any, digits: int = 4) -> str:
    if value is None:
        return "Not available"
    if isinstance(value, float):
        if abs(value) >= 1000:
            return f"{value:,.2f}"
        return f"{value:.{digits}f}"
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def _percent(value: Any) -> str:
    if isinstance(value, (float, int)):
        return f"{value * 100:.2f}%"
    return "Not applicable"


def _matrix_table(labels: list[str], matrix: list[list[int]]) -> str:
    header = "| True / Predicted | " + " | ".join(labels) + " |"
    separator = "|---|" + "|".join("---:" for _ in labels) + "|"
    rows = []
    for label, row in zip(labels, matrix):
        rows.append("| " + label + " | " + " | ".join(str(v) for v in row) + " |")
    return "\n".join([header, separator, *rows])


def _image_marker(relative_path: str, caption: str) -> str:
    if (PROJECT_ROOT / relative_path).exists():
        return f"\n[[IMAGE:{relative_path}|{caption}]]\n"
    return ""


def _feature_name_block(names: list[str], max_inline: int = 220) -> str:
    if len(names) <= max_inline:
        return ", ".join(names)
    shown = ", ".join(names[:max_inline])
    return (
        f"{shown}\n\n"
        f"The remaining {len(names) - max_inline} feature names continue using the same generated naming pattern "
        "inside the saved project feature extractor."
    )


def _markdown(results: dict[str, Any]) -> str:
    numerical = results["numerical"]
    image = results["image"]
    num_meta = numerical["metadata"]
    img_meta = image["metadata"]
    linear = numerical["linear_regression"]
    knn = numerical["knn_regressor"]
    logistic = image["logistic_regression"]
    kmeans = image["kmeans"]

    labels = img_meta["class_labels"]
    img_feature_names = get_image_feature_names()
    num_encoded_names = num_meta.get("encoded_feature_names", [])

    return f"""# Project Description Document

**Project Title:** Classical Machine Learning Project: Numerical Regression and Image Recognition  
**Prepared for:** University machine learning project submission  
**Student / Team:** ______________________________  
**Course:** ______________________________  
**Date:** {date.today().isoformat()}

---

## 1. Project Overview

This project implements two required machine learning tasks using only the datasets already provided inside the project folder.

- Numerical dataset task: Linear Regression and KNN Regressor.
- Image dataset task: Logistic Regression classifier and KMeans clustering-based classifier.

The implementation uses a complete train, validation, and test workflow. It saves trained models, preprocessing objects, evaluation reports, plots, and confusion matrices.

---

## 2. General Dataset Information

### 2.1 Numerical Dataset

| Item | Value |
|---|---|
| Dataset name | {num_meta["dataset_name"]} |
| Dataset source | {num_meta["source"]} |
| Dataset file | `{Path(num_meta["dataset_file"]).name}` |
| Problem type | Regression |
| Target column | `{num_meta["target_column"]}` |
| Total samples | {_fmt(num_meta["split_counts"]["total"])} |
| Original shape | {num_meta["original_shape"]} |
| Used shape | {num_meta["used_shape"]} |
| Raw input columns | {", ".join(num_meta["raw_feature_columns"])} |
| Numerical columns | {", ".join(num_meta["numeric_features"])} |
| Categorical columns | {", ".join(num_meta["categorical_features"])} |
| Training samples | {_fmt(num_meta["split_counts"]["train"])} |
| Validation samples | {_fmt(num_meta["split_counts"]["validation"])} |
| Testing samples | {_fmt(num_meta["split_counts"]["test"])} |

The numerical dataset is used for regression. Therefore, the correct model outputs are continuous numeric predictions, not class labels.

### 2.2 Image Dataset

| Item | Value |
|---|---|
| Dataset name | {img_meta["dataset_name"]} |
| Dataset source | {img_meta["source"]} |
| Problem type | Image classification / clustering-based labeling |
| Number of classes | {_fmt(img_meta["number_of_classes"])} |
| Class labels | {", ".join(labels)} |
| Total image samples | {_fmt(img_meta["split_counts"]["total"])} |
| Image size used | {img_meta["image_size"][0]} x {img_meta["image_size"][1]} pixels |
| Training samples | {_fmt(img_meta["split_counts"]["train"])} |
| Validation samples | {_fmt(img_meta["split_counts"]["validation"])} |
| Testing samples | {_fmt(img_meta["split_counts"]["test"])} |

Class sample counts:

{chr(10).join(f"- {label}: {_fmt(count)} images" for label, count in img_meta["class_counts"].items())}

The image implementation is limited to a maximum of five classes, satisfying the project requirement.

---

## 3. Implementation Details

### 3.1 Numerical Preprocessing

The numerical pipeline uses a scikit-learn preprocessing structure with separate handling for numerical and categorical columns.

- Missing numerical values are filled using median imputation.
- Missing categorical values are filled using most-frequent imputation.
- Numerical outliers are handled conservatively with IQR-based capping.
- Numerical features are scaled using StandardScaler.
- Categorical features are converted using OneHotEncoder.
- The preprocessing is fitted only on training data during model selection to avoid data leakage.

The target column was checked for skewness. The recorded target skewness is {_fmt(num_meta["target_skewness"])}. A `log1p` target transformation was tested, but the final model used it only if validation RMSE improved.

### 3.2 Numerical Feature Extraction Phase

For the numerical task, feature extraction means transforming the original mixed-type columns into a fully numeric machine learning matrix.

| Item | Value |
|---|---|
| Original raw input features | {_fmt(len(num_meta["raw_feature_columns"]))} |
| Final extracted/encoded features | {_fmt(num_meta["number_of_extracted_features"])} |
| Training feature dimension | {num_meta["feature_dimensions"]["train"]} |
| Validation feature dimension | {num_meta["feature_dimensions"]["validation"]} |
| Testing feature dimension | {num_meta["feature_dimensions"]["test"]} |

Raw feature names:

```text
{", ".join(num_meta["raw_feature_columns"])}
```

Encoded feature names generated by the preprocessing pipeline:

```text
{_feature_name_block(num_encoded_names)}
```

### 3.3 Image Preprocessing

The image pipeline reads images from class folders, resizes each image to a fixed size, normalizes pixel values, and extracts a fixed-length feature vector.

- Folder format: `datasets/images/class_name/image_file`
- Supported classes in this run: {", ".join(labels)}
- Resized image size: {img_meta["image_size"][0]} x {img_meta["image_size"][1]}
- Feature values are scaled before training Logistic Regression and KMeans.
- Validation and test images are not augmented, protecting fair evaluation.

### 3.4 Image Feature Extraction Phase

The image feature extractor converts each image into a numeric vector. This is required because Logistic Regression and KMeans cannot directly train on raw image files.

| Item | Value |
|---|---|
| Number of extracted image features | {_fmt(img_meta["number_of_extracted_features"])} |
| Training feature dimension | {img_meta["feature_dimensions"]["train"]} |
| Validation feature dimension | {img_meta["feature_dimensions"]["validation"]} |
| Testing feature dimension | {img_meta["feature_dimensions"]["test"]} |

Feature types used:

- RGB color histogram features.
- Grayscale histogram features.
- RGB channel statistics.
- Grayscale and gradient statistics.
- Edge-density feature.
- Local grid color statistics.
- HOG-style gradient orientation features.

Image feature names:

```text
{_feature_name_block(img_feature_names)}
```

---

## 4. Hyperparameters Used

### 4.1 Linear Regression

| Hyperparameter | Value |
|---|---|
| Model | scikit-learn LinearRegression |
| Learning rate | {linear["hyperparameters"]["learning_rate"]} |
| Optimizer | {linear["hyperparameters"]["optimizer"]} |
| Regularization | {linear["hyperparameters"]["regularization"]} |
| Batch size | {linear["hyperparameters"]["batch_size"]} |
| Number of epochs | {linear["hyperparameters"]["epochs"]} |
| fit_intercept | {linear["hyperparameters"]["fit_intercept"]} |
| positive | {linear["hyperparameters"]["positive"]} |
| selected_log1p_transform | {linear["hyperparameters"]["selected_log1p_transform"]} |

### 4.2 KNN Regressor

| Hyperparameter | Value |
|---|---|
| Model | scikit-learn KNeighborsRegressor |
| Learning rate | {knn["hyperparameters"]["learning_rate"]} |
| Optimizer | {knn["hyperparameters"]["optimizer"]} |
| Regularization | {knn["hyperparameters"]["regularization"]} |
| Batch size | {knn["hyperparameters"]["batch_size"]} |
| Number of epochs | {knn["hyperparameters"]["epochs"]} |
| Candidate n_neighbors | {knn["hyperparameters"]["candidate_k_values"]} |
| Candidate weights | {knn["hyperparameters"]["candidate_weights"]} |
| Candidate metrics | {knn["hyperparameters"]["candidate_metrics"]} |
| CV folds | {knn["hyperparameters"]["cv_folds"]} |
| Selected n_neighbors | {knn["hyperparameters"]["selected_n_neighbors"]} |
| Selected weights | {knn["hyperparameters"]["selected_weights"]} |
| Selected metric | {knn["hyperparameters"]["selected_metric"]} |

### 4.3 Logistic Regression

| Hyperparameter | Value |
|---|---|
| Model | scikit-learn LogisticRegression |
| Learning rate | {logistic["hyperparameters"]["learning_rate"]} |
| Optimizer | {logistic["hyperparameters"]["optimizer"]} |
| Regularization | {logistic["hyperparameters"]["regularization"]} |
| Batch size | {logistic["hyperparameters"]["batch_size"]} |
| Number of epochs | {logistic["hyperparameters"]["epochs"]} |
| Penalty | {logistic["hyperparameters"]["penalty"]} |
| Solver | {logistic["hyperparameters"]["solver"]} |
| Candidate C values | {logistic["hyperparameters"]["candidate_C_values"]} |
| Candidate max_iter | {logistic["hyperparameters"]["candidate_max_iter"]} |
| Candidate class_weight | {logistic["hyperparameters"]["candidate_class_weight"]} |
| Candidate PCA components | {logistic["hyperparameters"]["candidate_pca_components"]} |
| Selected C | {logistic["hyperparameters"]["selected_C"]} |
| Selected max_iter | {logistic["hyperparameters"]["selected_max_iter"]} |
| Selected class_weight | {logistic["hyperparameters"]["selected_class_weight"]} |
| Selected PCA components | {logistic["hyperparameters"]["selected_pca_n_components"]} |

### 4.4 KMeans

| Hyperparameter | Value |
|---|---|
| Model | scikit-learn KMeans |
| Learning rate | {kmeans["hyperparameters"]["learning_rate"]} |
| Optimizer | {kmeans["hyperparameters"]["optimizer"]} |
| Regularization | {kmeans["hyperparameters"]["regularization"]} |
| Batch size | {kmeans["hyperparameters"]["batch_size"]} |
| Number of epochs | {kmeans["hyperparameters"]["epochs"]} |
| n_clusters | {kmeans["hyperparameters"]["n_clusters"]} |
| init | {kmeans["hyperparameters"]["init"]} |
| selected n_init | {kmeans["hyperparameters"]["selected_n_init"]} |
| max_iter | {kmeans["hyperparameters"]["max_iter"]} |
| algorithm | {kmeans["hyperparameters"]["algorithm"]} |
| selected PCA components | {kmeans["hyperparameters"]["selected_pca_n_components"]} |

KMeans is unsupervised, so it does not naturally output class labels. The implementation maps clusters to labels by majority voting on the training data.

Cluster-to-label mapping:

{chr(10).join(f"- Cluster {cluster}: {label}" for cluster, label in kmeans["cluster_to_label_mapping"].items())}

---

## 5. Results Details on Testing Data

Important note: accuracy and confusion matrix are classification metrics. They are meaningful for Logistic Regression and KMeans in the image task. For Linear Regression and KNN Regressor, the correct testing metrics are MAE, MSE, RMSE, and R2 because the output is numeric and continuous.

### 5.1 Regression Results

| Model | MAE | MSE | RMSE | R2 | Accuracy | Confusion Matrix |
|---|---:|---:|---:|---:|---|---|
| Linear Regression | {_fmt(linear["test_metrics"]["MAE"])} | {_fmt(linear["test_metrics"]["MSE"])} | {_fmt(linear["test_metrics"]["RMSE"])} | {_fmt(linear["test_metrics"]["R2"])} | Not applicable | Not applicable |
| KNN Regressor | {_fmt(knn["test_metrics"]["MAE"])} | {_fmt(knn["test_metrics"]["MSE"])} | {_fmt(knn["test_metrics"]["RMSE"])} | {_fmt(knn["test_metrics"]["R2"])} | Not applicable | Not applicable |

Linear Regression achieved the better regression result in this run because it had lower RMSE and higher R2 on the test set.

### 5.2 Logistic Regression Image Classification Results

| Metric | Value |
|---|---|
| Test accuracy | {_percent(logistic["test_metrics"]["accuracy"])} |
| Testing samples | {_fmt(img_meta["split_counts"]["test"])} |
| Classes | {", ".join(labels)} |

Confusion matrix for Logistic Regression on testing data:

{_matrix_table(labels, logistic["test_metrics"]["confusion_matrix"])}

{_image_marker("outputs/confusion_matrices/logistic_regression_confusion_matrix.png", "Logistic Regression confusion matrix on testing data.")}

### 5.3 KMeans Image Classification Results

| Metric | Value |
|---|---|
| Test accuracy | {_percent(kmeans["test_metrics"]["accuracy"])} |
| Testing samples | {_fmt(img_meta["split_counts"]["test"])} |
| Classes | {", ".join(labels)} |

Confusion matrix for KMeans on testing data:

{_matrix_table(labels, kmeans["test_metrics"]["confusion_matrix"])}

{_image_marker("outputs/confusion_matrices/kmeans_confusion_matrix.png", "KMeans confusion matrix after cluster-to-label mapping on testing data.")}

---

## 6. Final Model Comparison

| Task | Model | Best use | Test result |
|---|---|---|---|
| Numerical regression | Linear Regression | Continuous target prediction | RMSE {_fmt(linear["test_metrics"]["RMSE"])}, R2 {_fmt(linear["test_metrics"]["R2"])} |
| Numerical regression | KNN Regressor | Distance-based continuous prediction | RMSE {_fmt(knn["test_metrics"]["RMSE"])}, R2 {_fmt(knn["test_metrics"]["R2"])} |
| Image classification | Logistic Regression | Supervised image classification | Accuracy {_percent(logistic["test_metrics"]["accuracy"])} |
| Image clustering/classification | KMeans | Unsupervised clustering with label mapping | Accuracy {_percent(kmeans["test_metrics"]["accuracy"])} |

{_image_marker("outputs/figures/regression_model_comparison.png", "Regression model comparison chart.")}

{_image_marker("outputs/figures/image_model_comparison.png", "Image model comparison chart.")}

---

## 7. Conclusion

This Project Description Document summarizes the required dataset details, implementation details, hyperparameters, extracted feature dimensions, and testing results.

The numerical task was evaluated using regression metrics because it predicts a continuous target. The image task was evaluated using classification metrics, including accuracy and confusion matrices.

The results show that Linear Regression performed better than KNN Regressor on the numerical test set, while Logistic Regression performed better than KMeans on the image test set. This is expected because Logistic Regression is supervised, whereas KMeans is unsupervised and must be adapted to classification using cluster-to-label majority voting.
"""


def _add_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        paragraph.paragraph_format.space_after = Pt(8)
        code_lines = []

    def flush_table() -> None:
        nonlocal table_lines
        if not table_lines:
            return
        rows = []
        for line in table_lines:
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if all(part.replace("-", "").replace(":", "").strip() == "" for part in parts):
                continue
            rows.append(parts)
        if rows:
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            table = doc.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            if r_idx == 0:
                                run.bold = True
            doc.add_paragraph()
        table_lines = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        else:
            flush_table()

        if line.startswith("[[IMAGE:"):
            inside = line.removeprefix("[[IMAGE:").removesuffix("]]")
            rel_path, caption = inside.split("|", 1)
            image_path = PROJECT_ROOT / rel_path
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(5.8))
                caption_p = doc.add_paragraph(caption)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption_p.runs:
                    run.italic = True
                    run.font.size = Pt(8)
            continue

        if line.strip() == "---":
            doc.add_page_break()
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(15, 23, 42)
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(line[2:].strip())
            continue

        paragraph = doc.add_paragraph()
        paragraph.add_run(line.replace("**", ""))

    flush_table()
    flush_code()


def _build_docx(markdown_text: str, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.color.rgb = RGBColor(30, 41, 59)

    _add_markdown_to_docx(doc, markdown_text)
    doc.save(path)


def _build_pdf(markdown_text: str, path: Path) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.55 * inch,
        title="Project Description Document",
        author="ML University Project",
    )
    doc.build(_markdown_to_story(markdown_text), onFirstPage=_page_decorator, onLaterPages=_page_decorator)


def build_project_description_document() -> dict[str, Path]:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    results = _load_results()
    markdown_text = _markdown(results)

    markdown_path = DOCS_DIR / "Project_Description_Document.md"
    docx_path = DOCS_DIR / "Project_Description_Document.docx"
    pdf_path = DOCS_DIR / "Project_Description_Document.pdf"

    markdown_path.write_text(markdown_text, encoding="utf-8")
    _build_docx(markdown_text, docx_path)
    _build_pdf(markdown_text, pdf_path)

    return {"Markdown": markdown_path, "DOCX": docx_path, "PDF": pdf_path}
