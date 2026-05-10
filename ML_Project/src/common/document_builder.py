from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List

from config import PROJECT_CONFIG


def _format_float(value: Any, decimals: int = 4) -> str:
    if isinstance(value, (int, float)):
        return f"{value:.{decimals}f}"
    return str(value)


def _format_money_metric(value: Any) -> str:
    if isinstance(value, (int, float)):
        return f"{value:,.2f}"
    return str(value)


def _classification_accuracy(value: Any) -> str:
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%"
    return str(value)


def build_cover_sheet_text() -> str:
    return f"""Project Cover Sheet

Faculty Name: ______________________________
Course Name: _______________________________
Project Title: Machine Learning Regression and Image Classification Project
Team Number: _______________________________

Team Members:
1. Name: ____________________________    ID: __________________
2. Name: ____________________________    ID: __________________
3. Name: ____________________________    ID: __________________
4. Name: ____________________________    ID: __________________

Submission Date: {date.today().isoformat()}

Project Summary:
This university machine learning project implements two regression models on a user-provided numerical dataset and two image classification/clustering models on a user-provided image dataset. The numerical task uses Linear Regression and KNN Regressor. The image task uses Logistic Regression and KMeans with cluster-to-label mapping by majority voting.
"""


def _metrics_table_markdown(title: str, metrics: Dict[str, Any], regression: bool = True) -> str:
    lines = [title, "", "| Metric | Value |", "|---|---|"]
    if regression:
        for metric in ["MAE", "MSE", "RMSE", "R2"]:
            value = metrics.get(metric, "Not available")
            formatter = _format_float if metric == "R2" else _format_money_metric
            lines.append(f"| {metric} | {formatter(value)} |")
    else:
        lines.append(f"| Accuracy | {_classification_accuracy(metrics.get('accuracy', 'Not available'))} |")
    return "\n".join(lines)


def build_project_description_text(results: Dict[str, Any]) -> str:
    numerical = results["numerical"]
    image = results["image"]
    numerical_meta = numerical["metadata"]
    image_meta = image["metadata"]

    linear_metrics = numerical["linear_regression"]["test_metrics"]
    knn_metrics = numerical["knn_regressor"]["test_metrics"]
    logistic_metrics = image["logistic_regression"]["test_metrics"]
    kmeans_metrics = image["kmeans"]["test_metrics"]
    numerical_target_transform = numerical_meta.get("final_target_transform", {})

    feature_names_numerical = numerical_meta["encoded_feature_names"]
    feature_names_image = image_meta["feature_names"]
    image_width, image_height = image_meta["image_size"]

    return f"""# Machine Learning Regression and Image Classification Project

## 1. Title Page

Project Title: Machine Learning Regression and Image Classification Project

Course: _______________________________

Faculty: ______________________________

Team Number: __________________________

Submission Date: {date.today().isoformat()}

## 2. Introduction

This project presents a complete supervised and unsupervised machine learning workflow using only the two datasets provided by the student. The first task is a numerical regression problem using a cyber security salary dataset. The second task is an image classification problem using a flower image dataset organized in class folders.

The project is designed for academic discussion. It separates preprocessing, feature engineering, training, evaluation, reporting, and saved model artifacts into clear Python modules. The code automatically inspects the provided datasets and uses one central configuration file for paths, target variables, class selection, image size, and model hyperparameters.

## 3. Objectives

- Build two regression models for a numerical dataset: Linear Regression and KNN Regressor.
- Build two image models with a maximum of five classes: Logistic Regression classifier and KMeans clustering-based classifier.
- Apply correct preprocessing for numerical and image datasets.
- Extract explainable features and document their dimensions.
- Evaluate regression using MAE, MSE, RMSE, and R2.
- Evaluate classification using accuracy, confusion matrix, and classification report.
- Save trained models, plots, reports, Word documents, and a PDF report.

## 4. Dataset 1: Numerical Dataset

Dataset name: {numerical_meta['dataset_name']}

Source: {numerical_meta['source']}

Dataset file: {numerical_meta['dataset_file']}

Total samples used: {numerical_meta['split_counts']['total']}

Original shape: {numerical_meta['original_shape']}

Used shape after removing missing target rows and configured leakage columns: {numerical_meta['used_shape']}

Target variable: {numerical_meta['target_column']}

Raw input columns: {', '.join(numerical_meta['raw_feature_columns'])}

Dropped columns: {', '.join(numerical_meta['dropped_columns']) if numerical_meta['dropped_columns'] else 'None'}

Missing value summary: {numerical_meta['missing_values']}

Target skewness before training: {_format_float(numerical_meta['target_skewness'])}

Target transform decision: {numerical_meta['target_transform_decision']['reason']}

Number of extracted/encoded features: {numerical_meta['number_of_extracted_features']}

Train/validation/test split:

- Training samples: {numerical_meta['split_counts']['train']}
- Validation samples: {numerical_meta['split_counts']['validation']}
- Testing samples: {numerical_meta['split_counts']['test']}

## 5. Dataset 2: Image Dataset

Dataset name: {image_meta['dataset_name']}

Source: {image_meta['source']}

Number of classes: {image_meta['number_of_classes']}

Class labels: {', '.join(image_meta['class_labels'])}

Samples per class: {image_meta['class_counts']}

Total image samples used: {image_meta['split_counts']['total']}

Image size used by the implementation: {image_width}x{image_height} pixels

Train/validation/test split:

- Training samples: {image_meta['split_counts']['train']}
- Validation samples: {image_meta['split_counts']['validation']}
- Testing samples: {image_meta['split_counts']['test']}

## 6. Preprocessing

For the numerical dataset, missing numerical values are imputed with the median, missing categorical values are imputed with the most frequent category, categorical features are one-hot encoded, and numerical values are standardized using StandardScaler. A conservative IQR-based capper is applied to numerical inputs before scaling to reduce the effect of extreme values without discarding many rows. The configured target column is excluded from the input matrix. The columns `salary` and `salary_currency` are removed because `salary_in_usd` is the selected target and these fields could create target leakage.

The target distribution is inspected before training. If the target is strongly skewed and a log transform improves validation RMSE, log1p is used through a transformed regressor. If it does not improve validation RMSE, the original target scale is kept.

For the image dataset, images are loaded from class folders under `datasets/images/`, converted to RGB, resized to {image_width}x{image_height}, normalized to the range 0 to 1, and split using stratification so that each class is represented in training, validation, and testing. Image feature scaling is handled with StandardScaler, and PCA is used only when it improves validation performance or training stability.

## 7. Feature Extraction

Numerical feature extraction:

- Original numerical columns: {', '.join(numerical_meta['numeric_features']) if numerical_meta['numeric_features'] else 'None'}
- Original categorical columns: {', '.join(numerical_meta['categorical_features']) if numerical_meta['categorical_features'] else 'None'}
- Encoding method: one-hot encoding for categorical variables
- Numerical outlier handling: IQR-based capping on the training split, then reuse of learned bounds on validation and test data
- Scaling method: StandardScaler for numerical variables
- Number of extracted/encoded features: {len(feature_names_numerical)}
- Final feature dimensions: train {numerical_meta['feature_dimensions']['train']}, validation {numerical_meta['feature_dimensions']['validation']}, test {numerical_meta['feature_dimensions']['test']}
- Feature names: {', '.join(feature_names_numerical)}

Image feature extraction:

- RGB color histograms: 16 bins per channel, producing 48 features
- Grayscale histogram: 32 bins
- RGB channel statistics: mean, standard deviation, minimum, and maximum for each channel, producing 12 features
- Grayscale and texture statistics: grayscale mean, grayscale standard deviation, gradient mean, gradient standard deviation, and edge density, producing 5 features
- Local color summary: a 4x4 spatial grid with mean and standard deviation for each RGB channel
- HOG-style gradient orientation features on grayscale image patches
- Total extracted image features: {len(feature_names_image)}
- Final feature dimensions: train {image_meta['feature_dimensions']['train']}, validation {image_meta['feature_dimensions']['validation']}, test {image_meta['feature_dimensions']['test']}
- Feature names: {', '.join(feature_names_image)}

## 8. Model 1: Linear Regression

Theory summary: Linear Regression estimates a continuous target value by learning a linear relationship between input features and the output. It is suitable as a clear baseline for regression problems.

Hyperparameters:

{numerical['linear_regression']['hyperparameters']}

Validation decision on target transform:

{numerical_target_transform}

Results on test data:

{_metrics_table_markdown('Linear Regression Test Metrics', linear_metrics)}

## 9. Model 2: KNN Regressor

Theory summary: KNN Regressor predicts a continuous value by finding the nearest training examples and averaging their target values. In this implementation, several k values are evaluated on the validation set, and the k value with the lowest validation RMSE is selected.

Hyperparameters:

{numerical['knn_regressor']['hyperparameters']}

Validation tuning results:

{numerical['knn_regressor']['validation_results_by_search']}

Results on test data:

{_metrics_table_markdown('KNN Regressor Test Metrics', knn_metrics)}

## 10. Model 3: Logistic Regression

Theory summary: Logistic Regression is a supervised classification model. For multiple classes, it learns decision boundaries that separate the class labels using the extracted image features. A small validation search is used to tune only C, max_iter, class_weight, and optional PCA components.

Hyperparameters:

{image['logistic_regression']['hyperparameters']}

Results on test data:

{_metrics_table_markdown('Logistic Regression Test Metrics', logistic_metrics, regression=False)}

Confusion matrix:

{logistic_metrics['confusion_matrix']}

## 11. Model 4: KMeans

Theory summary: KMeans is an unsupervised clustering algorithm that groups samples into clusters by minimizing within-cluster distance to centroids. Because KMeans does not naturally know class names, this project maps each cluster to a class label using majority voting on the training data.

Hyperparameters:

{image['kmeans']['hyperparameters']}

Cluster-to-label mapping method:

After fitting KMeans on the training features, each training image receives a cluster ID. For every cluster, the most frequent true class label among the training images in that cluster is assigned as the label for that cluster. Validation and test samples are then classified by predicting their cluster and replacing the cluster ID with the mapped class label.

Cluster-to-label mapping:

{image['kmeans']['cluster_to_label_mapping']}

Results on test data:

{_metrics_table_markdown('KMeans Test Metrics', kmeans_metrics, regression=False)}

Confusion matrix:

{kmeans_metrics['confusion_matrix']}

## 12. Comparison of Models

For regression, Linear Regression and KNN Regressor are compared using MAE, MSE, RMSE, and R2. Lower MAE, MSE, and RMSE indicate better numerical prediction error, while higher R2 indicates stronger explained variance.

For classification, Logistic Regression and KMeans are compared using accuracy and confusion matrices. Accuracy and confusion matrix are classification metrics only. They are not appropriate for regression because regression predicts continuous values rather than discrete class labels.

## 13. Discussion

The numerical task benefits from categorical encoding because salary datasets usually contain fields such as experience level, employment type, job title, employee residence, company location, and company size. Standardization helps KNN because distance-based models are sensitive to feature scale.

The improved numerical pipeline is stronger because preprocessing, clipping, and target-transform decisions are learned only from the training data. This reduces leakage, makes KNN less sensitive to outliers, and gives Linear Regression a fairer target distribution check without blindly forcing a log transform.

The image task uses compact, explainable features instead of a deep neural network. HOG-style grayscale features capture petal and edge structure, while color histograms and simple image statistics capture color cues that matter for flower classes. This keeps the method classical and discussion-friendly while improving generalization compared with raw pixels alone.

## 14. Challenges

- Avoiding target leakage in the salary regression task required removing columns that directly describe the same salary value in another format.
- KMeans required a careful cluster-to-label mapping step because it is an unsupervised method.
- Image classification needed a maximum of five classes, and the provided dataset already satisfies this requirement.
- Image features had to be simple enough for academic explanation while still useful for classification.
- KNN and Logistic Regression both required careful scaling so that one feature family would not dominate the distance or optimization process.

## 15. Conclusion

The project successfully implements four machine learning models across two different data types. Linear Regression and KNN Regressor are used for numerical regression, while Logistic Regression and KMeans are used for image classification and clustering-based classification. The project saves all models, metrics, plots, confusion matrices, and documents in a GitHub-ready structure.

## 16. References

- Scikit-learn documentation: LinearRegression, KNeighborsRegressor, LogisticRegression, KMeans, train_test_split, StandardScaler, OneHotEncoder.
- Pandas documentation for CSV loading and tabular data handling.
- NumPy documentation for numerical array operations.
- Matplotlib and Seaborn documentation for plotting.
- Pillow documentation for image loading and resizing.
"""


def _add_heading(document: Any, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def _is_table_separator(line: str) -> bool:
    return bool(line) and set(line) <= {"|", "-", " "}


def _parse_table_row(line: str) -> List[str]:
    return [cell.strip() for cell in line.strip("|").split("|")]


def _add_markdown_table(document: Any, table_lines: List[str]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    rows = [_parse_table_row(line) for line in table_lines if not _is_table_separator(line)]
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    if row_index == 0:
                        run.bold = True


def _add_paragraphs_from_markdown(document: Any, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            _add_heading(document, line[2:], 0)
        elif line.startswith("## "):
            _add_heading(document, line[3:], 1)
        elif line.startswith("### "):
            _add_heading(document, line[4:], 2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        else:
            document.add_paragraph(line)
        index += 1


def _build_docx(path: Path, title: str, body_text: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 22, RGBColor(31, 78, 121)),
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(67, 67, 67)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    document.core_properties.title = title
    _add_paragraphs_from_markdown(document, body_text)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _build_pdf(path: Path, title: str, body_text: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ProjectTitle",
            parent=styles["Title"],
            fontSize=20,
            textColor=colors.HexColor("#1F4E79"),
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading1"],
            fontSize=14,
            textColor=colors.HexColor("#1F4E79"),
            spaceBefore=12,
            spaceAfter=6,
        )
    )

    story: List[Any] = []
    for raw_line in body_text.splitlines():
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 0.08 * inch))
            continue
        if line.startswith("# "):
            story.append(Paragraph(line[2:], styles["ProjectTitle"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], styles["SectionHeading"]))
        elif line.startswith("- "):
            story.append(Paragraph(f"&bull; {line[2:]}", styles["BodyText"]))
        else:
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, styles["BodyText"]))

    path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=title,
    )
    document.build(story)


def build_project_documents(results: Dict[str, Any]) -> None:
    docs_dir = PROJECT_CONFIG["outputs"]["docs_dir"]
    cover_text = build_cover_sheet_text()
    description_text = build_project_description_text(results)

    (docs_dir / "Project_Cover_Sheet.md").write_text(cover_text, encoding="utf-8")
    (docs_dir / "Project_Description.md").write_text(description_text, encoding="utf-8")

    try:
        _build_docx(docs_dir / "Project_Cover_Sheet.docx", "Project Cover Sheet", cover_text)
        _build_docx(docs_dir / "Project_Description.docx", "Project Description", description_text)
        _build_pdf(docs_dir / "Project_Description.pdf", "Project Description", description_text)
    except ImportError as exc:
        warning = (
            "Document generation requires python-docx and reportlab. "
            "Install requirements.txt, then run main.py again. "
            f"Original error: {exc}"
        )
        (docs_dir / "DOCUMENT_GENERATION_WARNING.txt").write_text(warning, encoding="utf-8")
